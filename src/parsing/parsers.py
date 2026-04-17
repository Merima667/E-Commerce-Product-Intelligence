import csv
import xml.etree.ElementTree as ET
import os
import json
import sys
import pdfplumber
import re
from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook
from datetime import datetime

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'src')))
from storage.mongo import save_to_mongo 

raw_api_folder = "../../data/raw/api/"
raw_csv_folder = "../../data/raw/csv/"
raw_xml_folder = "../../data/raw/xml/"

def extract_review_fields(review):
    return {
        "review_id": review.get("review_id"),
        "title": review.get("review_title"),
        "comment": review.get("review_comment"),
        "rating": review.get("review_star_rating"),
        "author": review.get("review_author"),
        "author_id": review.get("review_author_id"),
        "date": review.get("review_date"),
        "verified": review.get("is_verified_purchase"),
        "helpful_votes": review.get("helpful_vote_statement"),
        "variant": review.get("reviewed_product_variant"),
        "review_link": review.get("review_link")
    }


def parse_json_files():
    for file_name in os.listdir(raw_api_folder):
        if file_name.endswith(".json"):
            file_path = os.path.join(raw_api_folder, file_name)
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            reviews = data.get("data", {}).get("reviews", [])
            for review in reviews:
                review_fields = extract_review_fields(review)
                print(f"Saving to MongoDB (JSON): {review_fields}")
                save_to_mongo(review_fields, file_name)

def parse_csv_file(csv_path):
    with open(csv_path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            review_fields = extract_review_fields(row)
            print(f"Saving to MongoDB (CSV): {review_fields}")
            save_to_mongo(review_fields, csv_path)

def parse_xml_file(xml_path):
    tree = ET.parse(xml_path)
    root = tree.getroot()
    for item in root.findall("review"):
        review = {
            "review_id": item.findtext("review_id"),
            "review_title": item.findtext("review_title"),
            "review_comment": item.findtext("review_comment"),
            "review_star_rating": item.findtext("review_star_rating"),
            "author_id": item.findtext("review_author"),
            "review_author_id": item.findtext("review_author_id"),
            "review_date": item.findtext("review_date"),
            "is_verified_purchase": item.findtext("is_verified_purchase") == "True",
            "helpful_vote_statement": item.findtext("helpful_vote_statement"),
            "reviewed_product_variant": item.findtext("reviewed_product_variant"),
            "review_link": item.findtext("review_link")
        }
        review_fields = extract_review_fields(review)
        print(f"Saving to MongoDB (XML): {review_fields}")
        save_to_mongo(review_fields, xml_path)

def _has_column_break(para):
    for run in para.runs:
        for br in run._r.findall(qn('w:br')):
            if br.get(qn('w:type')) == 'column':
                return True
    return False

def _get_num_columns(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        cols_el = sectPr.find(qn('w:cols'))
        if cols_el is not None:
            num = cols_el.get(qn('w:num'))
            if num and int(num) > 1:
                return int(num)
    return 1

def extract_text_from_pdf(pdf_path):
    pages = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = normalize_text(page.extract_text() or "")
            if text:
                pages.append(text)
    return "\n\n".join(pages)

def normalize_text(text):
    if not text:
        return ""
    
    text = re.sub(r"[\t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n\n", text)

    return text.strip()

import chardet

def read_file_with_encoding(file_path):
    with open(file_path, "rb") as f:
        raw = f.read()

    result = chardet.detect(raw)
    encoding = result.get("encoding") or "utf-8"
    confidence = result.get("confidence")

    print(f"Detected encoding: {encoding} (confidence: {confidence})")

    text = raw.decode(encoding, errors="replace")

    return text


def extract_text_from_two_column_pdf(pdf_path, gap=10):
    pages_text = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            mid_x = page.width / 2

            left_column = page.crop((0, 0, mid_x - gap, page.height))
            right_column = page.crop((mid_x + gap, 0, page.width, page.height))

            left_text = normalize_text(left_column.extract_text() or "")
            right_text = normalize_text(right_column.extract_text() or "")

            combined = "\n\n".join(part for part in [left_text, right_text] if part)
            if combined:
                pages_text.append(combined)

    return "\n\n".join(pages_text)

def extract_text_from_word(docx_path):
    doc = Document(docx_path)

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    tables_data = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                tables_data.append(" | ".join(row_data))
    return {
        "paragraphs": "\n\n".join(paragraphs),
        "tables": "\n".join(tables_data)
    }

def extract_text_from_two_column_word(docx_path):
    doc = Document(docx_path)
    num_cols = _get_num_columns(doc)

    if num_cols == 1:
        return extract_text_from_word(docx_path)
    
    columns = [[] for _ in range(num_cols)]
    current_col = 0

    for para in doc.paragraphs:
        if _has_column_break(para):
            current_col = min(current_col + 1, num_cols - 1)
            continue
        text = para.text.strip()
        if text:
            columns[current_col].append(text)
    tables_data = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                tables_data.append(" | ".join(row_data))
    
    result = {}
    for i, col_paragraphs in enumerate(columns):
        result[f"column_{i + 1}"] = "\n\n".join(col_paragraphs)
    result["tables"] = "\n".join(tables_data)
    result["full_text"] = "\n\n". join(normalize_text(col) for col in result.values() if isinstance(col,str))
    return result


def extract_data_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    ws = wb["Monthly Budget"]

    records = []

    for row in ws.iter_rows(min_row=4, values_only=True):
        if not row[0] or row[0] == "SUMMARY":
            continue
        record = {
            "category":     row[0],
            "description":  row[1],
            "budgeted":     row[2],
            "actual":       row[3],
            "variance":     row[4],
            "variance_pct": row[5],
            "status":       row[6],
            "notes":        row[7]
        }
        records.append(record)

    return records

def extract_transactions_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    ws = wb["Transaction Log"]
    transactions = []
    for row in ws.iter_rows(min_row=4, values_only=True):
        if not row[0]:
            continue
        transaction = {
            "date":            row[0],
            "description":     row[1],
            "category":        row[2],
            "amount":          row[3],
            "type":            row[4],
            "running_balance": row[5],
            "receipt":         row[6]
        }
        transactions.append(transaction)
    return transactions

def extract_savings_goals_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    ws = wb["Savings Goals"]
    goals = []
    for row in ws.iter_rows(min_row=3, values_only=True):
        if not row[0] or row[0] == "TOTAL":
            continue
        goal = {
            "goal_name": row[0],
            "target":    row[1],
            "saved":     row[2],
            "remaining": row[3],
            "progress":  row[4],
            "eta":       row[5]
        }
        goals.append(goal)
    return goals

def extract_trend_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    ws = wb["6-Month Trend"]
    trend = []
    for row in ws.iter_rows(min_row=3, values_only=True):
        if not row[0] or row[0] == "AVERAGE":
            continue
        trend.append({
            "month":        row[0],
            "income":       row[1],
            "expenses":     row[2],
            "savings":      row[3],
            "net_flow":     row[4],
            "savings_rate": row[5]
        })
    return trend

if __name__ == "__main__":
    parse_json_files()
    parse_csv_file(os.path.join(raw_csv_folder, "sample.csv"))
    parse_xml_file(os.path.join(raw_xml_folder, "sample.xml"))
    